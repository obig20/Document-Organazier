from pathlib import Path
from typing import List, Dict
from pydantic_settings import BaseSettings


class Settings(BaseSettings):
    """Application settings and constants."""

    # Server
    HOST: str = "127.0.0.1"
    PORT: int = 8000
    DEBUG: bool = True

    # Paths
    BASE_DIR: Path = Path(__file__).resolve().parents[2]
    DATA_DIR: Path = BASE_DIR / "data"
    UPLOAD_DIR: str = str(DATA_DIR / "uploads")
    STORAGE_DIR: str = str(DATA_DIR / "storage")

    # Database
    DATABASE_URL: str = f"sqlite:///{(BASE_DIR / 'data' / 'app.db').as_posix()}"

    # Uploads
    MAX_FILE_SIZE: int = 20 * 1024 * 1024  # 20 MB
    SUPPORTED_EXTENSIONS: set = {
        ".pdf", ".docx", ".txt", ".xlsx", ".pptx", ".jpg", ".jpeg", ".png", ".tiff", ".bmp"
    }

    # OCR / Tesseract
    TESSERACT_LANGUAGES: List[str] = ["eng", "amh"]
    TESSERACT_CONFIG: Dict[str, object] = {"config": "--oem 3 --psm 3", "timeout": 120}

    # Rule-based categories and keywords
    DOCUMENT_CATEGORIES: Dict[str, Dict[str, List[str]]] = {
        "demographics": {
            "keywords": [
                "name", "age", "gender", "address", "marital", "birth", "nationality",
                "እድሜ", "ጾታ", "ስም", "አድራሻ"
            ]
        },
        "housing": {
            "keywords": [
                "lease", "rent", "tenant", "landlord", "property", "housing", "contract",
                "ቤት", "ሕንጻ", "ኪራይ", "ኪራይ ስምምነት"
            ]
        },
        "id_registry": {
            "keywords": [
                "passport", "id", "identification", "registry", "certificate", "birth certificate",
                "መታወቂያ", "ፓስፖርት", "ማረጋገጫ"
            ]
        },
        "land_plans": {
            "keywords": [
                "land", "plot", "survey", "deed", "plan", "boundary",
                "መሬት", "እቃ", "እቅድ", "መለኪያ"
            ]
        },
        "other": {"keywords": []},
    }

    class Config:
        env_file = ".env"

    def ensure_directories(self) -> None:
        Path(self.UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
        Path(self.STORAGE_DIR).mkdir(parents=True, exist_ok=True)


settings = Settings()
settings.ensure_directories()
"""
Document processing module for AI Document Organizer
Handles OCR and text extraction from various file types
"""
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook
import magic
from pathlib import Path
from typing import Optional, Tuple
from datetime import datetime
import logging

from app.core.config import settings
from app.models.document import Document, DocumentMetadata

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        """Initialize the document processor"""
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract OCR is not installed or not in system PATH")
            raise
    
    def process_document(self, file_path: str) -> Tuple[Document, str]:
        """Process a document and extract its content"""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Initialize document
        doc = self._create_document(file_path)
        
        try:
            # Extract text based on file type
            file_ext = file_path.suffix.lower()
            
            # Check if file extension is supported
            if file_ext not in settings.SUPPORTED_EXTENSIONS:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Extract text based on file type
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_text_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                text = self._extract_text_from_excel(file_path)
            elif file_ext in ['.txt']:
                text = self._extract_text_from_txt(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                text = self._extract_text_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            doc.content = text
            doc.is_processed = True
            doc.processing_status = "completed"
            
        except Exception as e:
            error_msg = f"Error processing {file_path}: {str(e)}"
            logger.error(error_msg)
            doc.processing_status = "error"
            doc.error_message = error_msg
        
        return doc, doc.content
    
    def _create_document(self, file_path: Path) -> Document:
        """Create a Document instance with metadata"""
        file_stats = file_path.stat()
        return Document(
            filename=file_path.name,
            original_path=str(file_path.absolute()),
            stored_path="",
            title=file_path.stem,
            content="",
            category="other",
            confidence_score=0.0,
            metadata=DocumentMetadata(
                file_size=file_stats.st_size,
                file_type=file_path.suffix.lower(),
                mime_type=magic.from_file(str(file_path), mime=True),
                created_date=datetime.fromtimestamp(file_stats.st_ctime),
                modified_date=datetime.fromtimestamp(file_stats.st_mtime)
            )
        )
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF, with OCR fallback"""
        text_parts = []
        
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)
                    else:
                        pix = page.get_pixmap()
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        text = self._extract_text_from_image(img)
                        text_parts.append(text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.warning(f"PDF extraction failed, trying OCR fallback: {e}")
            return self._extract_text_from_image(file_path)
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        return "\n\n".join([p.text for p in doc.paragraphs])
    
    def _extract_text_from_excel(self, file_path: Path) -> str:
        """Extract text from Excel file"""
        wb = load_workbook(filename=file_path, read_only=True, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = []
            
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text:
                    sheet_text.append("\t".join(row_text))
            
            if sheet_text:
                text_parts.append(f"=== {sheet_name} ===\n" + "\n".join(sheet_text))
        
        return "\n\n".join(text_parts)
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _extract_text_from_image(self, image_source) -> str:
        """Extract text from image using Tesseract OCR"""
        if isinstance(image_source, (str, Path)):
            with Image.open(image_source) as img:
                return self._run_ocr(img)
        elif hasattr(image_source, 'save'):  # PIL Image
            return self._run_ocr(image_source)
        raise ValueError("Unsupported image source type")
    
    def _run_ocr(self, img: Image.Image) -> str:
        """Run Tesseract OCR on a PIL Image"""
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Use settings.TESSERACT_LANGUAGES and settings.TESSERACT_CONFIG
        languages = '+'.join(settings.TESSERACT_LANGUAGES)
        return pytesseract.image_to_string(
            img,
            lang=languages,
            config=settings.TESSERACT_CONFIG.get('config', '--oem 1 --psm 3'),
            timeout=settings.TESSERACT_CONFIG.get('timeout', 300)
        ).strip()

# Singleton instance
document_processor = DocumentProcessor()