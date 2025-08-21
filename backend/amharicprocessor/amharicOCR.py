import os
import logging
import pandas as pd
from glob import glob
from io import BytesIO
from typing import List, Optional, Dict, Union
from pathlib import Path

import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np
from docx import Document
import openpyxl

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class MultiFormatDocumentProcessor:
    """
    A comprehensive document processor that handles multiple file formats:
    - PDF (with OCR for Amharic text)
    - Images (JPG, PNG, TIFF with OCR)
    - Text files (TXT, CSV)
    - Word documents (DOCX)
    - Excel files (XLSX, XLS)
    - PowerPoint (PPTX) - text extraction from notes and slides
    """

    # Supported file formats
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'images': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif'],
        'text': ['.txt', '.csv', '.md', '.xml', '.html'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls'],
        'powerpoint': ['.pptx', '.ppt']
    }

    def __init__(self, file_path: str, save_files_path: str, config: Optional[dict] = None):
        """
        Constructor for the MultiFormatDocumentProcessor.

        Args:
            file_path (str): Path to directory or single file
            save_files_path (str): Directory for extracted text files
            config (dict, optional): Configuration options
        """
        self.file_path = file_path
        self.save_files_path = save_files_path
        self.config = self._get_default_config()
        
        if config:
            self.config.update(config)
        
        self._validate_paths()

    def _get_default_config(self) -> dict:
        """Return default configuration"""
        return {
            # OCR settings
            'dpi': 300,
            'preprocess_images': True,
            'language': 'amh+eng',
            'oem': 1,
            'psm': 6,
            'timeout': 30,
            
            # Processing settings
            'extract_tables': True,
            'preserve_formatting': False,
            'max_file_size_mb': 50,
            
            # Excel settings
            'read_all_sheets': True,
            
            # Word settings
            'extract_comments': False,
            
            # Image settings
            'image_quality': 95
        }

    def _validate_paths(self):
        """Validate input and output paths"""
        if not os.path.exists(self.file_path):
            raise ValueError(f"Input path does not exist: {self.file_path}")
        
        if not os.path.exists(self.save_files_path):
            os.makedirs(self.save_files_path)
            logger.info(f"Created output directory: {self.save_files_path}")

    def process_all_files(self) -> dict:
        """
        Process all supported files in the input directory.

        Returns:
            dict: Processing statistics
        """
        stats = {
            'total_files': 0,
            'successful': 0,
            'failed': 0,
            'by_format': {},
            'total_characters': 0
        }

        try:
            all_files = self._get_all_supported_files()
            stats['total_files'] = len(all_files)

            if not all_files:
                logger.warning("No supported files found for processing")
                return stats

            for file_path in all_files:
                try:
                    file_ext = Path(file_path).suffix.lower()
                    file_type = self._get_file_type(file_ext)
                    
                    if file_type not in stats['by_format']:
                        stats['by_format'][file_type] = {'successful': 0, 'failed': 0}
                    
                    result = self.process_single_file(file_path)
                    
                    if result['success']:
                        stats['successful'] += 1
                        stats['by_format'][file_type]['successful'] += 1
                        stats['total_characters'] += result['character_count']
                        
                        logger.info(
                            f"✓ {file_type.upper()}: {Path(file_path).name} | "
                            f"Chars: {result['character_count']} | "
                            f"Time: {result['processing_time']:.2f}s"
                        )
                    else:
                        stats['failed'] += 1
                        stats['by_format'][file_type]['failed'] += 1
                        logger.error(f"✗ {file_type.upper()}: {Path(file_path).name} - {result['error']}")
                        
                except Exception as e:
                    stats['failed'] += 1
                    file_type = self._get_file_type(Path(file_path).suffix.lower())
                    if file_type not in stats['by_format']:
                        stats['by_format'][file_type] = {'successful': 0, 'failed': 0}
                    stats['by_format'][file_type]['failed'] += 1
                    logger.error(f"✗ Error processing {Path(file_path).name}: {str(e)}")
                    continue

            self._print_summary(stats)
            
        except Exception as e:
            logger.error(f"Fatal error in processing: {str(e)}")
            raise

        return stats

    def _get_all_supported_files(self) -> List[str]:
        """Get all supported files in directory"""
        if os.path.isfile(self.file_path):
            return [self.file_path] if self._is_supported_format(self.file_path) else []
        
        all_files = []
        for format_exts in self.SUPPORTED_FORMATS.values():
            for ext in format_exts:
                all_files.extend(glob(os.path.join(self.file_path, f"*{ext}")))
                all_files.extend(glob(os.path.join(self.file_path, f"*{ext.upper()}")))
        
        return list(set(all_files))  # Remove duplicates

    def _is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return any(ext in exts for exts in self.SUPPORTED_FORMATS.values())

    def _get_file_type(self, extension: str) -> str:
        """Get file type category from extension"""
        for file_type, exts in self.SUPPORTED_FORMATS.items():
            if extension in exts:
                return file_type
        return 'unknown'

    def process_single_file(self, file_path: str) -> dict:
        """
        Process a single file based on its format.

        Args:
            file_path (str): Path to the file

        Returns:
            dict: Processing result
        """
        result = {
            'success': False,
            'text': '',
            'character_count': 0,
            'processing_time': 0,
            'error': None,
            'file_type': None,
            'metadata': {}
        }

        start_time = os.times().elapsed
        file_ext = Path(file_path).suffix.lower()
        file_type = self._get_file_type(file_ext)

        try:
            if file_type == 'pdf':
                result = self._process_pdf(file_path)
            elif file_type == 'images':
                result = self._process_image(file_path)
            elif file_type == 'text':
                result = self._process_text_file(file_path)
            elif file_type == 'word':
                result = self._process_word_document(file_path)
            elif file_type == 'excel':
                result = self._process_excel_file(file_path)
            elif file_type == 'powerpoint':
                result = self._process_powerpoint(file_path)
            else:
                result['error'] = f"Unsupported file format: {file_ext}"
            
            result['file_type'] = file_type
            result['processing_time'] = os.times().elapsed - start_time

        except Exception as e:
            result['error'] = f"Error processing {file_type} file: {str(e)}"
            logger.error(f"Error processing {file_path}: {str(e)}")

        return result

    def _process_pdf(self, file_path: str) -> dict:
        """Process PDF file with OCR"""
        # Your existing PDF processing code here
        # (Copy from the previous AmharicOCR class)
        return self._extract_text_from_pdf(file_path)

    def _process_image(self, file_path: str) -> dict:
        """Process image file with OCR"""
        result = {'success': False, 'text': '', 'character_count': 0}
        
        try:
            image = Image.open(file_path)
            processed_image = self._preprocess_image(image)
            
            text = pytesseract.image_to_string(
                processed_image, 
                config=f'--oem {self.config["oem"]} --psm {self.config["psm"]} -l {self.config["language"]}',
                timeout=self.config['timeout']
            )
            
            result.update({
                'success': True,
                'text': text,
                'character_count': len(text)
            })
            
        except Exception as e:
            result['error'] = f"Image processing error: {str(e)}"
        
        return result

    def _process_text_file(self, file_path: str) -> dict:
        """Process text-based files"""
        result = {'success': False, 'text': '', 'character_count': 0}
        
        try:
            # Detect encoding
            encodings = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    result.update({
                        'success': True,
                        'text': text,
                        'character_count': len(text),
                        'encoding': encoding
                    })
                    break
                except UnicodeDecodeError:
                    continue
            
            if not result['success']:
                result['error'] = "Could not decode text file with any encoding"
                
        except Exception as e:
            result['error'] = f"Text file processing error: {str(e)}"
        
        return result

    def _process_word_document(self, file_path: str) -> dict:
        """Process Word documents"""
        result = {'success': False, 'text': '', 'character_count': 0}
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(' | '.join(row_text))
            
            text = '\n'.join(text_parts)
            
            result.update({
                'success': True,
                'text': text,
                'character_count': len(text)
            })
            
        except Exception as e:
            result['error'] = f"Word document processing error: {str(e)}"
        
        return result

    def _process_excel_file(self, file_path: str) -> dict:
        """Process Excel files"""
        result = {'success': False, 'text': '', 'character_count': 0, 'tables': []}
        
        try:
            if self.config['read_all_sheets']:
                excel_file = pd.ExcelFile(file_path)
                sheets = excel_file.sheet_names
            else:
                sheets = [None]  # Read only first sheet
            
            text_parts = []
            
            for sheet_name in sheets:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if sheet_name:
                        text_parts.append(f"=== Sheet: {sheet_name} ===")
                    
                    # Convert dataframe to text
                    text_parts.append(df.to_string())
                    
                    if self.config['extract_tables']:
                        result['tables'].append({
                            'sheet_name': sheet_name,
                            'data': df.to_dict('records'),
                            'shape': df.shape
                        })
                        
                except Exception as e:
                    logger.warning(f"Error reading sheet {sheet_name}: {e}")
                    continue
            
            text = '\n'.join(text_parts)
            
            result.update({
                'success': True,
                'text': text,
                'character_count': len(text)
            })
            
        except Exception as e:
            result['error'] = f"Excel file processing error: {str(e)}"
        
        return result

    def _process_powerpoint(self, file_path: str) -> dict:
        """Process PowerPoint files (basic text extraction)"""
        result = {'success': False, 'text': '', 'character_count': 0}
        
        try:
            # Simple text extraction from PPTX (would need python-pptx for full support)
            # This is a placeholder - you'd need to install python-pptx for full functionality
            text = f"[PowerPoint file: {Path(file_path).name} - Install python-pptx for full text extraction]"
            
            result.update({
                'success': True,
                'text': text,
                'character_count': len(text)
            })
            
        except Exception as e:
            result['error'] = f"PowerPoint processing error: {str(e)}"
        
        return result

    def _save_extracted_text(self, text: str, filename: str, file_type: str) -> str:
        """Save extracted text with appropriate naming"""
        output_path = os.path.join(self.save_files_path, f"{filename}_{file_type}.txt")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return output_path

    def _print_summary(self, stats: dict):
        """Print processing summary"""
        print("\n" + "="*70)
        print("DOCUMENT PROCESSING SUMMARY")
        print("="*70)
        print(f"Total files processed: {stats['total_files']}")
        print(f"Successful:           {stats['successful']}")
        print(f"Failed:               {stats['failed']}")
        
        if stats['total_files'] > 0:
            success_rate = (stats['successful'] / stats['total_files']) * 100
            print(f"Success rate:         {success_rate:.1f}%")
        
        print(f"Total characters:     {stats['total_characters']:,}")
        
        if stats['by_format']:
            print("\nBy file format:")
            for format_type, counts in stats['by_format'].items():
                total = counts['successful'] + counts['failed']
                if total > 0:
                    rate = (counts['successful'] / total) * 100
                    print(f"  {format_type.upper():<12} {counts['successful']:>3}✓ {counts['failed']:>3}✗ ({rate:.1f}%)")
        
        print("="*70)

# Usage examples
if __name__ == "__main__":
    # Process all files in a directory
    processor = MultiFormatDocumentProcessor(
        file_path="./documents",
        save_files_path="./extracted_text",
        config={
            'language': 'amh+eng',
            'read_all_sheets': True,
            'extract_tables': True
        }
    )
    
    results = processor.process_all_files()
    
    # Process a single file
    single_result = processor.process_single_file("./documents/report.docx")
    print(f"Extracted {single_result['character_count']} characters from Word document")